import os
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = upload_dir
        os.makedirs(upload_dir, exist_ok=True)
        self._init_parsers()

    def _init_parsers(self):
        try:
            import pdfplumber
            self.pdf_parser = pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed, PDF parsing may be limited")
            self.pdf_parser = None

        try:
            from docx import Document
            self.docx_parser = Document
        except ImportError:
            logger.warning("python-docx not installed, Word parsing may be limited")
            self.docx_parser = None

        try:
            import markdown
            self.md_parser = markdown
        except ImportError:
            logger.warning("markdown not installed, Markdown parsing may be limited")
            self.md_parser = None

        try:
            import pytesseract
            from PIL import Image
            self.ocr = pytesseract
            self.image = Image
        except ImportError:
            logger.warning("pytesseract or Pillow not installed, OCR will be unavailable")
            self.ocr = None
            self.image = None

    def get_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'word',
            '.doc': 'word',
            '.md': 'markdown',
            '.markdown': 'markdown',
            '.txt': 'text',
            '.png': 'image',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.gif': 'image',
            '.bmp': 'image',
            '.tiff': 'image',
        }
        return type_map.get(ext, 'unknown')

    def parse(self, file_path: str, file_type: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        parsers = {
            'pdf': self._parse_pdf,
            'word': self._parse_word,
            'markdown': self._parse_markdown,
            'text': self._parse_text,
            'image': self._parse_image,
        }

        parser = parsers.get(file_type)
        if not parser:
            raise ValueError(f"Unsupported file type: {file_type}")

        return parser(file_path)

    def _parse_pdf(self, file_path: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        if not self.pdf_parser:
            raise ImportError("pdfplumber is required for PDF parsing")

        content_parts = []
        page_contents = []
        page_count = 0

        with self.pdf_parser.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                tables = page.extract_tables()

                page_content = text
                if tables:
                    for table in tables:
                        table_text = self._table_to_markdown(table)
                        page_content += f"\n\n{table_text}"

                content_parts.append(page_content)
                page_contents.append({
                    'page_number': page_num,
                    'content': page_content,
                    'has_tables': len(tables) > 0
                })

        full_content = "\n\n".join(content_parts)
        return full_content, page_count, page_contents

    def _table_to_markdown(self, table: List[List[Any]]) -> str:
        if not table or not table[0]:
            return ""

        def clean_cell(cell):
            if cell is None:
                return ""
            return str(cell).strip().replace('\n', ' ')

        rows = []
        header = "| " + " | ".join(clean_cell(cell) for cell in table[0]) + " |"
        separator = "| " + " | ".join("---" for _ in table[0]) + " |"
        rows.append(header)
        rows.append(separator)

        for row in table[1:]:
            row_text = "| " + " | ".join(clean_cell(cell) for cell in row) + " |"
            rows.append(row_text)

        return "\n".join(rows)

    def _parse_word(self, file_path: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        if not self.docx_parser:
            raise ImportError("python-docx is required for Word document parsing")

        doc = self.docx_parser(file_path)
        content_parts = []
        page_contents = []

        for para in doc.paragraphs:
            if para.text.strip():
                content_parts.append(para.text.strip())

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            if table_data:
                content_parts.append(self._table_to_markdown(table_data))

        full_content = "\n\n".join(content_parts)
        page_count = max(1, len(content_parts) // 10)
        page_contents.append({
            'page_number': 1,
            'content': full_content,
            'has_tables': len(doc.tables) > 0
        })

        return full_content, page_count, page_contents

    def _parse_markdown(self, file_path: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        page_count = max(1, content.count('\n# ') + 1)
        page_contents = [{
            'page_number': 1,
            'content': content,
            'has_tables': '|' in content and '---' in content
        }]

        return content, page_count, page_contents

    def _parse_text(self, file_path: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()

        lines = content.split('\n')
        page_count = max(1, len(lines) // 50)
        page_contents = [{
            'page_number': 1,
            'content': content,
            'has_tables': False
        }]

        return content, page_count, page_contents

    def _parse_image(self, file_path: str) -> Tuple[str, int, List[Dict[str, Any]]]:
        if not self.ocr or not self.image:
            raise ImportError("pytesseract and Pillow are required for image OCR")

        try:
            img = self.image.open(file_path)
            text = self.ocr.image_to_string(img, lang='chi_sim+eng')
            page_contents = [{
                'page_number': 1,
                'content': text.strip(),
                'has_tables': False,
                'is_ocr': True
            }]
            return text.strip(), 1, page_contents
        except Exception as e:
            logger.error(f"OCR failed for {file_path}: {e}")
            raise
